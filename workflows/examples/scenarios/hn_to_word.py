"""场景2: 访问 Hacker News，抓取前10条新闻标题保存为Word文档."""

import asyncio
from pathlib import Path
from datetime import datetime
from omniauto.core.context import TaskContext, StepResult
from omniauto.core.state_machine import AtomicStep, Workflow
from docx import Document

async def hn_to_word(ctx: TaskContext) -> StepResult:
    browser = ctx.browser_state.get("browser")
    if browser is None:
        raise RuntimeError("浏览器引擎未初始化")

    # 1. 访问 HN
    await browser.goto("https://news.ycombinator.com")
    await asyncio.sleep(3)

    # 2. 提取标题和链接
    page = browser.page
    items = await page.eval_on_selector_all(
        '.titleline > a',
        '''elements => elements.slice(0, 10).map(e => ({title: e.innerText, href: e.href}))'''
    )
    items = [i for i in (items or []) if i and i.get("title") and i.get("title").strip()]

    # 3. 保存Word
    output_dir = Path("./outputs")
    output_dir.mkdir(exist_ok=True)
    word_path = output_dir / f"hn_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    doc = Document()
    doc.add_heading("Hacker News 每日精选", level=1)
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. ").bold = True
        p.add_run(item["title"])
        p.add_run(f"\n   链接: {item['href']}")

    doc.save(str(word_path))

    return StepResult(success=True, data={"word_path": str(word_path), "count": len(items)})

workflow = Workflow(task_id="scenario_hn_to_word")
workflow.add_step(AtomicStep("fetch_and_save_word", hn_to_word, lambda r: r.success))
